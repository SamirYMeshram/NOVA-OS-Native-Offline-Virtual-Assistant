from __future__ import annotations
from pathlib import Path
from .base import Extractor
from nova.documents.models import Document

class DocxExtractor(Extractor):
    extensions = {'.docx'}

    def extract(self, path: Path) -> Document:
        try:
            import docx  # type: ignore
            doc = docx.Document(str(path))
            text = '\n'.join(p.text for p in doc.paragraphs)
        except Exception as exc:
            text = f"[DOCX extraction unavailable: install python-docx. Error: {exc}]"
        return Document(path=path, text=text, metadata={'type': 'docx'})
