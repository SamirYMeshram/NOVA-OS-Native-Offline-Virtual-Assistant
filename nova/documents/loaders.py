from __future__ import annotations

import csv
import json
from dataclasses import dataclass
from pathlib import Path

from nova.core.config import SecurityConfig
from nova.core.security import SafetyGuard

SUPPORTED_EXTENSIONS = {
    ".txt", ".md", ".markdown", ".py", ".js", ".ts", ".tsx", ".jsx", ".json", ".csv",
    ".log", ".yaml", ".yml", ".toml", ".ini", ".html", ".css", ".sql", ".pdf", ".docx", ".xlsx"
}


@dataclass(slots=True)
class LoadedDocument:
    path: Path
    text: str
    metadata: dict[str, str]


class DocumentLoader:
    def __init__(self, security: SecurityConfig | None = None) -> None:
        self.security = security or SecurityConfig()
        self.guard = SafetyGuard(self.security)

    def can_load(self, path: Path) -> bool:
        return path.suffix.lower() in SUPPORTED_EXTENSIONS

    def load(self, path: str | Path) -> LoadedDocument:
        file_path = Path(path).expanduser().resolve()
        decision = self.guard.check_path_read(file_path)
        if not decision.allowed:
            raise PermissionError(decision.reason)
        if file_path.stat().st_size > self.security.max_file_read_mb * 1024 * 1024:
            raise ValueError(f"File is larger than configured read limit: {file_path}")
        ext = file_path.suffix.lower()
        if ext == ".pdf":
            text = self._load_pdf(file_path)
        elif ext == ".docx":
            text = self._load_docx(file_path)
        elif ext == ".xlsx":
            text = self._load_xlsx(file_path)
        elif ext == ".csv":
            text = self._load_csv(file_path)
        elif ext == ".json":
            text = self._load_json(file_path)
        else:
            text = file_path.read_text(encoding="utf-8", errors="ignore")
        return LoadedDocument(
            path=file_path,
            text=self.guard.redact_secrets(text),
            metadata={"extension": ext, "size": str(file_path.stat().st_size)},
        )

    @staticmethod
    def _load_pdf(path: Path) -> str:
        try:
            from pypdf import PdfReader  # type: ignore
        except Exception as exc:
            raise RuntimeError("PDF support requires: pip install pypdf") from exc
        reader = PdfReader(str(path))
        return "\n\n".join(page.extract_text() or "" for page in reader.pages)

    @staticmethod
    def _load_docx(path: Path) -> str:
        try:
            import docx  # type: ignore
        except Exception as exc:
            raise RuntimeError("DOCX support requires: pip install python-docx") from exc
        doc = docx.Document(str(path))
        return "\n".join(p.text for p in doc.paragraphs)

    @staticmethod
    def _load_xlsx(path: Path) -> str:
        try:
            from openpyxl import load_workbook  # type: ignore
        except Exception as exc:
            raise RuntimeError("Excel support requires: pip install openpyxl") from exc
        wb = load_workbook(path, read_only=True, data_only=True)
        lines: list[str] = []
        for sheet in wb.worksheets:
            lines.append(f"# Sheet: {sheet.title}")
            for row in sheet.iter_rows(values_only=True):
                lines.append("\t".join("" if v is None else str(v) for v in row))
        return "\n".join(lines)

    @staticmethod
    def _load_csv(path: Path) -> str:
        lines: list[str] = []
        with path.open("r", encoding="utf-8", errors="ignore", newline="") as handle:
            reader = csv.reader(handle)
            for idx, row in enumerate(reader):
                lines.append("\t".join(row))
                if idx > 10_000:
                    lines.append("[truncated after 10000 rows]")
                    break
        return "\n".join(lines)

    @staticmethod
    def _load_json(path: Path) -> str:
        data = json.loads(path.read_text(encoding="utf-8", errors="ignore"))
        return json.dumps(data, indent=2, ensure_ascii=False)
